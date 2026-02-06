import re
import os
import uuid
import io
from docx import Document
from minio import Minio
from dotenv import load_dotenv

load_dotenv()

# MinIO config
mc = Minio(
    endpoint=os.getenv("MINIO_ENDPOINT"),
    access_key=os.getenv("MINIO_ACCESS_KEY"),
    secret_key=os.getenv("MINIO_SECRET_KEY"),
    secure=os.getenv("MINIO_SECURE", "False").lower() == "true"
)

BK = os.getenv("MINIO_BUCKET", "ctu-quyche")
IBK = os.getenv("MINIO_IMAGE_BUCKET", "ctu-quyche-images")

# Tạo bucket nếu chưa có
for b in [BK, IBK]:
    if not mc.bucket_exists(b):
        mc.make_bucket(b)

def clean_text(txt: str) -> str:
    return re.sub(r'\s+', ' ', txt).strip()

def upload_image(ib: bytes, fn: str) -> str:
    on = f"images/{fn}"
    mc.put_object(
        bucket_name=IBK,
        object_name=on,
        data=io.BytesIO(ib),
        length=len(ib),
        content_type="image/png"
    )
    prot = "https" if os.getenv("MINIO_SECURE", "False").lower() == "true" else "http"
    return f"{prot}://{os.getenv('MINIO_ENDPOINT')}/{IBK}/{on}"

def table_to_md(tbl) -> str:
    # Chuyển bảng Word sang markdown
    if not tbl.rows:
        return ""
    
    hd = [clean_text(c.text) for c in tbl.rows[0].cells]
    md = "| " + " | ".join(hd) + " |\n"
    md += "| " + "--- | " * len(hd) + "\n"
    
    for r in tbl.rows[1:]:
        cs = [clean_text(c.text) for c in r.cells]
        md += "| " + " | ".join(cs) + " |\n"
    
    return md.strip()

def extract_sections_and_images(fp: str, mw: int = 400, mtr: int = 20):
    doc = Document(fp)
    cks = []
    cur = {
        "chương": "", "điều": "", "khoản": "", "điểm": "",
        "lines": [], "images": [], "full_titles": []
    }

    # Trích xuất ảnh
    ic = 0
    ius = []
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            ic += 1
            blob = rel.target_part.blob
            ext = "png" if b"png" in blob[:8] else "jpg"
            base = os.path.basename(fp).replace(".docx", "")
            fn = f"{base}_img_{ic:03d}.{ext}"
            url = upload_image(blob, fn)
            ius.append(url)

    ii = 0

    # Iterate qua paragraphs và tables
    for el in doc.element.body:
        if el.tag.endswith('p'):  # Paragraph
            para = next(p for p in doc.paragraphs if p._element == el)
            txt = para.text.strip()
            if not txt and not para._element.xpath('.//w:drawing'):
                continue

            # Kiểm tra có ảnh
            hd = bool(para._element.xpath('.//w:drawing'))
            if hd and ii < len(ius):
                cur["images"].append(ius[ii])
                ii += 1

            if not txt:
                continue

            # Phát hiện tiêu đề
            if re.match(r'^CHƯƠNG\s+[IVXLD]+', txt, re.I):
                _save_ck(cur, cks, fp, mw)
                t = re.sub(r'^CHƯƠNG\s+[IVXLD]+\.?\s*-?\s*', '', txt, flags=re.I).strip()
                cur = {"chương": t or "Không có tên chương", "điều": "", "khoản": "", "điểm": "",
                       "lines": [txt], "images": cur["images"][-1:] if hd else [], "full_titles": [txt]}

            elif re.match(r'^Điều\s+\d+', txt, re.I):
                _save_ck(cur, cks, fp, mw)
                m = re.search(r'Điều\s+(\d+)', txt, re.I)
                num = m.group(1) if m else ""
                t = re.sub(r'^Điều\s+\d+\.?\s*-?\s*', '', txt, flags=re.I).strip()
                cur = {"chương": cur["chương"], "điều": t or f"Điều {num}", "khoản": "", "điểm": "",
                       "lines": [txt], "images": cur["images"][-1:] if hd else [], "full_titles": [txt]}

            elif re.match(r'^\d+\.', txt) and len(txt.split()) <= 15:
                _save_ck(cur, cks, fp, mw)
                cur = {**cur, "khoản": txt, "điểm": "", "lines": [txt], "images": cur["images"][-1:] if hd else []}

            elif re.match(r'^[a-z]\)|^\d+\)', txt):
                cur["điểm"] = txt
                cur["lines"].append(txt)

            else:
                cur["lines"].append(txt)

        elif el.tag.endswith('tbl'):  # Table
            tbl = next(t for t in doc.tables if t._element == el)
            tmd = table_to_md(tbl)
            if tmd:
                tw = len(tmd.split())
                if tw <= mw and len(tbl.rows) <= mtr:
                    cur["lines"].append("Bảng dữ liệu:\n" + tmd)
                else:
                    # Bảng dài: chia chunk
                    _save_ck(cur, cks, fp, mw)
                    hd = [clean_text(c.text) for c in tbl.rows[0].cells]
                    rs = [[clean_text(c.text) for c in r.cells] for r in tbl.rows[1:]]
                    
                    for i in range(0, len(rs), mtr):
                        crs = rs[i:i + mtr]
                        cmd = "| " + " | ".join(hd) + " |\n| " + "--- | " * len(hd) + "\n"
                        for r in crs:
                            cmd += "| " + " | ".join(r) + " |\n"
                        
                        sfx = f" (phần {i//mtr + 1})" if len(rs) > mtr else ""
                        tck = {
                            "chương": cur["chương"],
                            "điều": cur["điều"],
                            "khoản": cur["khoản"],
                            "điểm": cur["điểm"],
                            "lines": ["Bảng dữ liệu" + sfx + ":\n" + cmd.strip()],
                            "images": [],
                            "full_titles": cur["full_titles"]
                        }
                        _save_ck(tck, cks, fp, mw)
                    cur = {"chương": cur["chương"], "điều": cur["điều"], "khoản": cur["khoản"], "điểm": "",
                           "lines": [], "images": [], "full_titles": cur["full_titles"]}

    _save_ck(cur, cks, fp, mw)
    return cks

def _save_ck(cur, cks, sf, mw):
    if not cur["lines"]:
        return
    ct = "\n".join(cur["lines"])
    ws = ct.split()

    tp = [p for p in [cur["chương"], cur["điều"], cur["khoản"], cur["điểm"]] if p]
    tit = " > ".join(tp) if tp else "Nội dung chung"

    imgs = cur["images"]

    if len(ws) <= mw:
        cks.append(_make_ck(ct, tit, imgs, sf))
    else:
        st = mw - 60
        for i in range(0, len(ws), st):
            cht = " ".join(ws[i:i+mw])
            sfx = "" if i == 0 else " (tiếp theo)"
            cis = imgs if i == 0 else []
            cks.append(_make_ck(cht, tit + sfx, cis, sf))

def _make_ck(ct, tit, imgs, src):
    return {
        "content": ct,
        "images": imgs,
        "metadata": {
            "title": tit,
            "level": len(tit.split(" > ")),
            "source": os.path.basename(src),
            "chunk_id": str(uuid.uuid4())
        }
    }